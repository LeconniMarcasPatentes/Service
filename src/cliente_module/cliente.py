from planilha_module import Planilha
from utils import *
import pandas as pd
from docx import Document
from datetime import datetime

# Classe Cliente
class Cliente:
    def __init__(self, razao_social, pf_pj, site_empresa, cpf_cnpj, tipo_empresa, endereco, telefones_empresa, 
                nome_contato, email_contato, cargo_contato, telefones_contato, apresentacao
                ):
        self.razao_social      = razao_social       # Razão Social/Nome
        self.pf_pj             = pf_pj              # PF/PJ
        self.site_empresa      = site_empresa       # Endereço do Site da Empresa
        self.cpf_cnpj          = cpf_cnpj           # CPF/CNPJ
        self.tipo_empresa      = tipo_empresa       # Tipo da Empresa (PJ, MEI, etc.)
        self.endereco          = endereco           # Endereço da Empresa (Dicionário)
        self.telefones_empresa = telefones_empresa  # Telefones da Empresa (Lista)
        self.nome_contato      = nome_contato       # Nome do Contato
        self.email_contato     = email_contato      # E-mail do Contato
        self.cargo_contato     = cargo_contato      # Cargo do Contato
        self.telefones_contato = telefones_contato  # Telefones de Contatos (Lista)
        self.apresentacao      = apresentacao       # Tipo de Apresentação (Figurativa, Mista, Nominativa)
    # __init__ ( )
    
    """ Cliente toString """
    def __str__( self ):
        return (f"Cliente:\n"
                f"Razão Social: {self.razao_social}\n" 
                f"CPF/CNPJ....: {self.cpf_cnpj}\n"
                f"Endereço....: {self.endereco}\n"
                f"Porte.......: {self.tipo_empresa}\n"
                )
    # __str__ ( )
    
    """ Atualizar a planilha de clientes, adicionando novos clientes ou modificando existentes. """
    def atualizar_planilha_clientes( clientes ):
        output_filepath = 'data/processed/Clientes.xlsx'

        try:
            try:
                df_existente = pd.read_excel( output_filepath )
            except FileNotFoundError:
                df_existente = pd.DataFrame( )

            novos_clientes = []

            for res in clientes:
                cliente_data = {
                    'Razao Social'      : res.cliente_razao_social,
                    'PF/PJ'             : res.cliente_pf_pj,
                    'CPF/CNPJ'          : res.cliente_cpf_cpnj,
                    'Tipo Empresa'      : res.cliente_tipo,
                    'Logradouro'        : res.cliente_endereco['logradouro'],
                    'Numero'            : res.cliente_endereco['numero'],
                    'Bairro'            : res.cliente_endereco['bairro'],
                    'Complemento'       : res.cliente_endereco['complemento'],
                    'Municipio'         : res.cliente_endereco['municipio'],
                    'UF'                : res.cliente_endereco['uf'],
                    'CEP'               : res.cliente_endereco['cep'],
                    'Telefone Empresa 1': res.cliente_telefone[0] if len(res.cliente_telefone) > 0 else '',
                    'Telefone Empresa 2': res.cliente_telefone[1] if len(res.cliente_telefone) > 1 else '',
                    'Telefone Empresa 3': res.cliente_telefone[2] if len(res.cliente_telefone) > 2 else '',
                    'Site Empresa'      : res.cliente_site,
                    'Nome Contato'      : res.contato_nome,
                    'Email Contato'     : res.contato_email,
                    'Cargo Contato'     : res.contato_cargo,
                    'Telefone Contato 1': res.contato_telefone[0] if len(res.contato_telefone) > 0 else '',
                    'Telefone Contato 2': res.contato_telefone[1] if len(res.contato_telefone) > 1 else '',
                    'Telefone Contato 3': res.contato_telefone[2] if len(res.contato_telefone) > 2 else '',
                    'Apresentacao'      : res.cliente_apresentacao
                }

                # Verificar se o cliente já existe na planilha (baseado em CPF/CNPJ)
                cliente_existente = df_existente[df_existente['CPF/CNPJ'] == cliente_data['CPF/CNPJ']]

                if cliente_existente.empty:
                    novos_clientes.append(cliente_data)
                else:
                    df_existente.update(pd.DataFrame([cliente_data]))

            if novos_clientes:
                df_novos = pd.DataFrame(novos_clientes)
                df_atualizado = pd.concat([df_existente, df_novos], ignore_index=True)
            else:
                df_atualizado = df_existente

            df_atualizado.to_excel(output_filepath, index=False)

            print( f"Planilha de clientes atualizada com sucesso: {output_filepath}" )

        except Exception as e:
            print( f"Erro ao atualizar a planilha de clientes: {e}" )
    # end atualizar_planilha_clientes ( )

    """ Escrever os dados de um cliente no contrato. """
    def escrever_contrato( tipo_contrato, cliente ):
        filepath1 = ""
        filepath2 = ""
        outputpath = ""
        if tipo_contrato == "":
            contrato1( filepath1, outputpath, cliente )
        elif tipo_contrato == "":
            contrato2( filepath2, outputpath, cliente )
    # end escrever_contrato ( )
# Clientes

def contrato1( template_path, output_path, cliente ):
    doc = Document( template_path )
    
    data_atual = datetime.now().strftime("%d/%m/%Y")  # Formato: dia/mês/ano
    
    for paragrafo in doc.paragraphs :
        texto = paragrafo.text
        # texto = texto.split()
        
        if '[NOME_CONTRATANTE],' in texto:
            paragrafo.text = paragrafo.text.replace( '[NOME_CONTRATANTE],', cliente.razao_social.upper( )+ "," )
            
        if '[NACIONALIDADE_CONTRATANTE]' in texto:
            paragrafo.text = paragrafo.text.replace( '[NACIONALIDADE_CONTRATANTE]', 'brasileira' )
            
        if '[CNPJ_CONTRATANTE],' in texto:
            paragrafo.text = paragrafo.text.replace( '[CNPJ_CONTRATANTE],', formatar_cpf(cliente.cpf_cnpj)+"," if cliente == 'PF' else formatar_cnpj(cliente.cpf_cnpj)+"," ) 
                
        if '[ENDERCO_CONTRATANTE],' in texto:
            endereco_completo = f"{cliente.endereco['logradouro']}, {cliente.endereco['numero']}, {cliente.endereco['bairro']}, {cliente.endereco['municipio']}/{cliente.endereco['uf']} - {formatar_cep(cliente.endereco['cep'])}"  
            paragrafo.text = paragrafo.text.replace('[ENDERCO_CONTRATANTE],', endereco_completo+",")
                        
        if '[TITULO_PATENTE]' in texto:
            paragrafo.text = paragrafo.text.replace( '[TITULO_PATENTE]', cliente.razao_social.upper( ) )   
        
        if '[VALOR_SERVIÇO]' in texto:
            paragrafo.text = paragrafo.text.replace( '[VALOR_SERVIÇO]', '10.000,00')  
                
        if '[VALOR_INICIAL]' in texto:
            paragrafo.text = paragrafo.text.replace( '[VALOR_INICIAL]', '5.000,00') 
            
        if '[VALOR_FINAL]' in texto:
            paragrafo.text = paragrafo.text.replace( '[VALOR_FINAL]', '5.000,00')  
                
        if '[CIDADE_CONTRATANTE]' in texto:
            paragrafo.text = paragrafo.text.replace( '[CIDADE_CONTRATANTE]', cliente.endereco['municipio'] )
                    
        if '[LOCAL_DATA]' in texto:
            paragrafo.text = paragrafo.text.replace( '[LOCAL_DATA]', data_atual )

    doc.save( output_path )
# end contrato1 ( )

def contrato2( template_path, output_path, cliente ):
    doc = Document( template_path )
    
    data_atual = datetime.now().strftime("%d/%m/%Y")  # Formato: dia/mês/ano
    
    for paragrafo in doc.paragraphs :
        texto = paragrafo.text
        # texto = texto.split()
        
        if '[NOME_CONTRATANTE],' in texto:
            paragrafo.text = paragrafo.text.replace( '[NOME_CONTRATANTE],', cliente.razao_social.upper( )+ "," )
            
        if '[NACIONALIDADE_CONTRATANTE]' in texto:
            paragrafo.text = paragrafo.text.replace( '[NACIONALIDADE_CONTRATANTE]', 'brasileira' )
            
        if '[CNPJ_CONTRATANTE],' in texto:
            paragrafo.text = paragrafo.text.replace( '[CNPJ_CONTRATANTE],', formatar_cpf(cliente.cpf_cnpj)+"," if cliente == 'PF' else formatar_cnpj(cliente.cpf_cnpj)+"," ) 
                
        if '[ENDERCO_CONTRATANTE],' in texto:
            endereco_completo = f"{cliente.endereco['logradouro']}, {cliente.endereco['numero']}, {cliente.endereco['bairro']}, {cliente.endereco['municipio']}/{cliente.endereco['uf']} - {formatar_cep(cliente.endereco['cep'])}"  
            paragrafo.text = paragrafo.text.replace('[ENDERCO_CONTRATANTE],', endereco_completo+",")
                        
        if '[TITULO_PATENTE]' in texto:
            paragrafo.text = paragrafo.text.replace( '[TITULO_PATENTE]', cliente.razao_social.upper( ) )   
        
        if '[VALOR_SERVIÇO]' in texto:
            paragrafo.text = paragrafo.text.replace( '[VALOR_SERVIÇO]', '10.000,00')  
                
        if '[VALOR_INICIAL]' in texto:
            paragrafo.text = paragrafo.text.replace( '[VALOR_INICIAL]', '5.000,00') 
            
        if '[VALOR_FINAL]' in texto:
            paragrafo.text = paragrafo.text.replace( '[VALOR_FINAL]', '5.000,00')  
                
        if '[CIDADE_CONTRATANTE]' in texto:
            paragrafo.text = paragrafo.text.replace( '[CIDADE_CONTRATANTE]', cliente.endereco['municipio'] )
                    
        if '[LOCAL_DATA]' in texto:
            paragrafo.text = paragrafo.text.replace( '[LOCAL_DATA]', data_atual )

    doc.save( output_path )
# end contrato2 ( )
